from __future__ import annotations

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


SECTION_15_HEADING = "15 OVERVIEW OF SIGNALS: NEW, ONGOING, OR CLOSED"
SECTION_15_1_SUBHEADING = "15.1 BRIEF DESCRIPTION OF SIGNAL DETECTION METHOD"
SECTION_15_2_SUBHEADING = "15.2 SIGNAL TABULATIONS"

SECTION_15_1_TEXT_LINES = [
    "Signal detection involves generation of frequency report from company safety database, scientific literature review, excluded cases, potential risks of the product, and safety-related updates from PRAC recommendations.",
    "The quantitative method was applied on events (identified from frequency report) to identify the relevant DECs for further evaluation from company safety database.",
]

THRESHOLDS = [
    "Nij >= 3 during the reporting period",
    "Nij >= 3 during the cumulative period",
]

SECTION_15_POST_THRESHOLD_TEXT = [
    "The events crossing the threshold as per defined parameter were considered as DECs.",
    "The DECs were assessed for labelling as per RSI. If the DEC is labeled as per RSI, frequency trend was evaluated.",
    "If DEC is labeled/covered under medical concept based on the information available in innovator/ comparator’s RSI and unlabeled as per MAH RSI, recommendation for label update is proposed.",
    "If the DEC is unlabeled as per innovator/ comparator and MAH RSI, the DEC was taken up further for the clinical assessment. The clinical evaluation included the assessment of the causal association between the drug and event and characterization of the qualified signal based on the cumulative cases of the event in company safety database.",
    "The identified signals were further categorized as validated, non-validated signals and events under monitoring signal bases on evidence of strength. Signals with insufficient or inconclusive information and signals which can be explained on the basis of presence of confounding factors such as presence of multiple suspect drugs or concomitant medications or significant medical history were categorized as Non-validated signals.",
    "The signals with sufficient evidences demonstrating the existence of a new potentially causal association or a new aspect of a known association and therefore justifies further analysis of the signal were categorized as validated signals.",
    "Signals with equivocal information which cannot be categorized neither non-validated nor has sufficient evidence to be considered as validated, these were categorized as events under monitoring signals and kept for further monitoring.",
]


def _add_table_from_dataframe(doc: Document, df: pd.DataFrame, title: str | None = None) -> None:
    if title:
        p = doc.add_paragraph(title)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if df is None or df.empty:
        doc.add_paragraph("No data available")
        return
    table = doc.add_table(rows=1, cols=len(df.columns), style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(df.columns):
        hdr_cells[i].text = str(h)
    for _, row in df.iterrows():
        tr = table.add_row().cells
        for i, cell_val in enumerate(row):
            tr[i].text = str(cell_val)


def write_section_15(doc: Document, *, table_section15: pd.DataFrame, closed_signal: list[str]) -> None:
    doc.add_heading(SECTION_15_HEADING, level=1)

    doc.add_heading(SECTION_15_1_SUBHEADING, level=2)
    for line in SECTION_15_1_TEXT_LINES:
        doc.add_paragraph(line)
    for item in THRESHOLDS:
        doc.add_paragraph(item, style="List Bullet")
    # Add requested follow-up narrative immediately after thresholds
    for line in SECTION_15_POST_THRESHOLD_TEXT:
        doc.add_paragraph(line)

    doc.add_paragraph("")
    doc.add_heading(SECTION_15_2_SUBHEADING, level=2)

    signals_count = 0 if table_section15 is None or table_section15.empty else len(table_section15)
    title = (
        f"There were {signals_count} signals for the product identified during the period covered by this report."
    )
    _add_table_from_dataframe(doc, table_section15, title=title)

