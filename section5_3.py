from __future__ import annotations

import re
from typing import NamedTuple, Optional, Sequence, List
from dataclasses import dataclass

import numpy as np
import pandas as pd
import requests
import urllib3
from bs4 import BeautifulSoup
from docx import Document


class ExposureComputationResult(NamedTuple):
    country_table: pd.DataFrame
    non_country_table: pd.DataFrame
    country_total: int
    non_country_total: int
    combined_total: int
    ddd_value: Optional[float]


def extract_table_after_text(doc) -> Optional[List[List[str]]]:
    """Find a marker text and return the first table after it as list-of-lists."""
    search_text = "Cumulative sales data sale required"
    pattern = re.compile(re.escape(search_text), re.IGNORECASE)
    found_index = None

    for i, para in enumerate(doc.paragraphs):
        if pattern.search(para.text):
            found_index = i
            break
    if found_index is None:
        return None

    para_counter = 0
    table_counter = 0
    for block in doc.element.body:
        if block.tag.endswith('p'):
            para_counter += 1
        elif block.tag.endswith('tbl'):
            if para_counter > found_index:
                table = doc.tables[table_counter]
                break
            table_counter += 1
    else:
        return None

    table_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    if len(table_data) > 1 and table_data[0] == table_data[1]:
        table_data.pop(1)
    return table_data


def _create_clean_total_row(dataframe: pd.DataFrame, total_col: str = "Patients Exposure (PTY) for period") -> pd.DataFrame:
    total = dataframe[total_col].sum(numeric_only=True)
    total_row = {col: "" for col in dataframe.columns}
    total_row["Country"] = "Total"
    total_row[total_col] = int(total)
    return pd.DataFrame([total_row])


def _map_dosage(product_name: str, product_dosage_map: dict) -> str:
    if pd.isna(product_name):
        return ""
    name = str(product_name).lower()
    for key, val in product_dosage_map.items():
        if key.lower() in name:
            return val
    return ""


def _add_dosage_column(df: pd.DataFrame, product_dosage_map: dict) -> pd.DataFrame:
    col_to_use = next((c for c in ["Product", "Molecule"] if c in df.columns), None)
    if col_to_use is None:
        return df
    df["Dosage Form (Units)"] = df[col_to_use].apply(_map_dosage, args=(product_dosage_map,))
    return df


def prepare_exposure_tables(
    dataframe: pd.DataFrame,
    ddd_value: Optional[float | int],
    country_name: str,
    country_aliases: Optional[Sequence[str]] = None,
) -> ExposureComputationResult:
    """Calculate exposure metrics from a DataFrame for section 5.3."""

    product_dosage_map = {
        "Esomeprazole": "Gastro-resistant",
        "JUBIGORD 20": "Gastro-resistant",
        "JUBIGORD 40": "Gastro-resistant",
        "Esomeprazol": "Gastro-resistant",
        "JUBIUM": "Gastro-resistant",
        "Zipola 5": "Film coated Tablet",
        "Zipola 10": "Film coated Tablet",
        "Jubilonz OD10": "Oro dispersible tablet",
        "Jubilonz OD5": "Oro dispersible tablet",
        "SCHIZOLANZ": "Oro dispersible tablet",
        "Olanzapine film coated tablets": "Film coated Tablet",
        "Olanzapine": "Film coated Tablet",
    }

    dataframe = _add_dosage_column(dataframe, product_dosage_map)
    dataframe = dataframe.drop_duplicates().reset_index(drop=True)

    if "Strength in mg" in dataframe.columns:
        dataframe["Strength in mg"] = pd.to_numeric(
            dataframe["Strength in mg"].astype(str).str.replace("mg", "", regex=False).str.strip(), errors="coerce"
        )

    pack_column = next((col for col in ("Pack", "Packs") if col in dataframe.columns), None)
    if pack_column:
        dataframe[pack_column] = pd.to_numeric(
            dataframe[pack_column].astype(str).str.replace(",", "", regex=False).str.extract(r"(\d+)")[0], errors="coerce"
        ).fillna(0).astype(int)
        dataframe.drop(columns=[pack_column], inplace=True)

    if "Pack size" in dataframe.columns:
        pack_size_series = dataframe["Pack size"].astype(str).str.split(":", n=1).str[-1].str.strip()
        extracted = pack_size_series.str.extract(r"(\d+)\s*[xX]\s*(\d+)")
        dataframe["Pack size"] = (
            pd.to_numeric(extracted[0], errors="coerce").fillna(1).astype(int) *
            pd.to_numeric(extracted[1], errors="coerce").fillna(1).astype(int)
        )

    unit_column = "Number of tablets / Capsules/Injections"
    if unit_column in dataframe.columns:
        dataframe[unit_column] = pd.to_numeric(
            dataframe[unit_column].astype(str).str.replace(",", "", regex=False).str.split(":").str[-1].str.strip(),
            errors="coerce",
        )

    if "Delivered quantity (mg)" in dataframe.columns:
        dataframe.drop(columns=["Delivered quantity (mg)"], inplace=True)

    if "Product" in dataframe.columns and "Molecule" not in dataframe.columns:
        dataframe.rename(columns={"Product": "Molecule"}, inplace=True)

    normalized_ddd = float(ddd_value) if ddd_value not in (None, "") and not pd.isna(ddd_value) else None
    dataframe["DDD*"] = f"{int(normalized_ddd)} mg" if normalized_ddd else ""

    if unit_column in dataframe.columns and "Strength in mg" in dataframe.columns:
        dataframe["Sales Figure (mg) or period/Volume of sales (in mg)"] = dataframe[unit_column] * dataframe["Strength in mg"]
    else:
        dataframe["Sales Figure (mg) or period/Volume of sales (in mg)"] = np.nan

    dataframe["Patients Exposure (PTY) for period"] = (
        (dataframe["Sales Figure (mg) or period/Volume of sales (in mg)"] / (normalized_ddd * 365)) if normalized_ddd else np.nan
    )
    dataframe["Patients Exposure (PTY) for period"] = dataframe["Patients Exposure (PTY) for period"].round(0)

    if "Country" not in dataframe.columns:
        dataframe["Country"] = "Unknown"

    dataframe.fillna("", inplace=True)

    alias_candidates: List[str] = list(country_aliases) if country_aliases else []
    alias_candidates.append(country_name)
    target_countries = {str(val).strip().casefold() for val in alias_candidates if val}

    if (country_name or "").strip().casefold() == "eu&uk":
        eu_uk_countries = {"uk", "se", "dk"}
        country_mask = dataframe["Country"].astype(str).str.strip().str.casefold().isin(eu_uk_countries)
    else:
        country_mask = dataframe["Country"].astype(str).str.strip().str.casefold().isin(target_countries)

    df_country = dataframe[country_mask].copy()
    df_non_country = dataframe[~country_mask].copy()

    total_column = "Patients Exposure (PTY) for period"
    df_country_total = _create_clean_total_row(df_country, total_column)
    df_non_country_total = _create_clean_total_row(df_non_country, total_column)

    df_country = pd.concat([df_country, df_country_total], ignore_index=True)
    df_non_country = pd.concat([df_non_country, df_non_country_total], ignore_index=True)

    country_total = int(df_country_total[total_column].iloc[0]) if not df_country_total.empty else 0
    non_country_total = int(df_non_country_total[total_column].iloc[0]) if not df_non_country_total.empty else 0
    combined_total = country_total + non_country_total

    final_column_order = [
        "Country", "Molecule", "Dosage Form (Units)", "Pack size", "DDD*",
        "Sales Figure (mg) or period/Volume of sales (in mg)", "Patients Exposure (PTY) for period",
    ]

    def reorder_columns(df, ordered_cols):
        new_order = [col for col in ordered_cols if col in df.columns]
        remaining_cols = [col for col in df.columns if col not in new_order]
        return df[new_order + remaining_cols]

    df_country = reorder_columns(df_country, final_column_order)
    df_non_country = reorder_columns(df_non_country, final_column_order)

    return ExposureComputationResult(
        country_table=df_country,
        non_country_table=df_non_country,
        country_total=country_total,
        non_country_total=non_country_total,
        combined_total=combined_total,
        ddd_value=float(normalized_ddd) if normalized_ddd is not None else None,
    )


# DDD fallback (remote source)
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)


def fetch_ddd_fallback(medicine: str, code: str | float | int | None) -> float:
    """Fetch a DDD value from a public site as last resort.

    Returns numpy.nan when not found or on error.
    """
    if code is None or (isinstance(code, float) and np.isnan(code)):
        return np.nan
    url = f"https://atcddd.fhi.no/atc_ddd_index/?code={code}"
    try:
        response = requests.get(url, verify=False, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, "html.parser")
        for td in soup.find_all("td", align="right"):
            value = td.get_text(strip=True)
            if value.replace('.', '', 1).isdigit():
                return float(value)
        return np.nan
    except Exception:
        return np.nan


# -----------------------
# Accumulation utilities
# -----------------------

@dataclass
class Section5_3Data:
    medicine: str
    place: str
    date: str
    country_name: str
    results: Optional[ExposureComputationResult]


def accumulate_section5_3(*, docx_path: str, ddd_excel_path: str | None,
                          country: str, medicine: str, place: str, date: str) -> Section5_3Data:
    try:
        ddd_df = pd.read_excel(ddd_excel_path) if ddd_excel_path else pd.DataFrame()
    except Exception:
        ddd_df = pd.DataFrame()

    if not docx_path:
        return Section5_3Data(medicine=medicine, place=place, date=date, country_name=country, results=None)

    doc = Document(docx_path)
    table_data = extract_table_after_text(doc)

    # DDD Value
    ddd_value = np.nan
    ddd_row = pd.DataFrame()
    if not ddd_df.empty and "Drug Name" in ddd_df.columns:
        ddd_row = ddd_df[ddd_df["Drug Name"].astype(str).str.lower() == medicine.lower()]
    if not ddd_row.empty and "DDD Value" in ddd_row.columns and pd.notna(ddd_row.iloc[0]["DDD Value"]):
        ddd_value = ddd_row.iloc[0]["DDD Value"]
    else:
        code = ddd_row.iloc[0]["Drug Code"] if not ddd_row.empty and "Drug Code" in ddd_row.columns else np.nan
        ddd_value = fetch_ddd_fallback(medicine, code)

    results: Optional[ExposureComputationResult] = None
    if table_data and pd.notna(ddd_value):
        sales_df = pd.DataFrame(table_data[1:], columns=table_data[0])
        computed = prepare_exposure_tables(
            dataframe=sales_df, ddd_value=ddd_value, country_name=country,
        )
        results = computed

    return Section5_3Data(
        medicine=medicine,
        place=place,
        date=date,
        country_name=country,
        results=results,
    )


