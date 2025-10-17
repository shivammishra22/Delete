from __future__ import annotations
from os.path import abspath, dirname, join
import sys
sys.path.insert(0, abspath(join(dirname(__file__), '..')))
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from add_cover_page import PSURGenerator
from config.styling_config import DocumentStyling

from extractors.section5_2 import accumulate_section5_2
from extractors.section5_3 import accumulate_section5_3
from extractors.section6_3 import accumulate_section6_3
from extractors.section15 import accumulate_section15

from writers.section5_2 import write_section_5_2
from writers.section5_3 import generate_exposure_document, generate_fallback_doc
from writers.section6_3 import write_section_6_3
from writers.section15 import write_section_15


# --------------------
# Simple configuration
# --------------------

# Fill these paths for your dataset
SECTION5_2_DRA_PATH = r"C:\Users\shivam.mishra2\Downloads\ALL_PSUR_File\PSUR_all _Data\Levetiracetam PSUR (Ukraine)_30-Nov-2021 to 30-Nov-2024\Draft\Clinical study & BA BE\Data request form.docx"  # e.g. r"C:\\path\\to\\Data request form.docx"
REPORTING_PERIOD = "30-Nov-2021 to 30-Nov-2024"  # e.g. "30-Nov-2021 to 30-Nov-2024"

SECTION5_3_DOCX_PATH =  r"C:\Users\shivam.mishra2\Downloads\ALL_PSUR_File\PSUR_all _Data\Esomeprazole ADCO_19-Apr-2017 to 31-Dec-2024\Draft\DRA\Data request form_Esomeprazole ADCO.docx"  # Word doc with cumulative sales data
DDD_EXCEL_PATH = r"C:\Users\shivam.mishra2\Downloads\FINAL_PSUR_REPORT\drug_code_map_Section_5.xlsx"      # Excel with Drug Name / DDD Value / Drug Code
COUNTRY = "eu&uk"
MEDICINE = "Esomeprazole"
PLACE = "EU & UK"
DATE = "17=07-2020"

CUM_EXCEL_PATH = r"C:\Users\shivam.mishra2\Downloads\ALL_PSUR_File\PSUR_all _Data\Levetiracetam PSUR (Ukraine)_30-Nov-2021 to 30-Nov-2024\Draft\ICH LL\Levetiracetam_Draft\Levetiracetam_DRAFT Cumulative ICH LL.xlsx"  # cumulative excel/csv
CUM_RTF_PATH =  r"C:\Users\shivam.mishra2\Downloads\Levetiracetam_Cumulative period ICH LL.rtf"   # cumulative rtf
INT_EXCEL_PATH = r"C:\Users\shivam.mishra2\Downloads\ALL_PSUR_File\PSUR_all _Data\Levetiracetam PSUR (Ukraine)_30-Nov-2021 to 30-Nov-2024\Draft\ICH LL\Levetiracetam_Final\Levetiracetam_ICH LL_Reporting Period_Final.xlsx"  # optional interval excel/csv
INT_RTF_PATH = r"C:\Users\shivam.mishra2\Downloads\ALL_PSUR_File\PSUR_all _Data\Levetiracetam PSUR (Ukraine)_30-Nov-2021 to 30-Nov-2024\Draft\ICH LL\Levetiracetam_Draft\Levetiracetam UA_PSUR_Draft Reporting Period.rtf"

SECTION15_INPUT_PATH = r"C:\Users\shivam.mishra2\Downloads\New_Psur_File\psur_iteration2\Data request form_RK.DOCX",

# Cover page / branding
LOGO_PATH = r"C:\Users\shivam.mishra2\Downloads\New_Psur_File\jub.png"
DRUG_NAME = MEDICINE
REPORTING_PERIOD_COVER = REPORTING_PERIOD
VERSION_STATUS = "Draft"
VERSION_DATE = DATE or "22-12-2025"
DRUG_CODE = "ACO222"


def add_cover_and_toc(doc: Document, *, logo_path: str | None, title: str, subtitle: str) -> None:
    """Add a simple cover page and a TOC placeholder to the document."""
    # Cover page
    p = doc.add_paragraph("")
    run = p.add_run(title)
    run.font.name = DocumentStyling.FONT_NAME
    run.font.size = Pt(24)
    run.bold = True
    p.alignment = DocumentStyling.CENTER_ALIGNMENT

    p2 = doc.add_paragraph(subtitle)
    p2.alignment = DocumentStyling.CENTER_ALIGNMENT
    for r in p2.runs:
        r.font.name = DocumentStyling.FONT_NAME
        r.font.size = Pt(12)

    # Spacer
    doc.add_paragraph("")

    # Table of Contents placeholder (Word needs update fields to populate)
    doc.add_paragraph("Table of Contents").style = 'Heading 1'
    paragraph = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    # Escape backslashes so Python parser does not treat \u as unicode escape
    fld.set(qn('w:instr'), 'TOC \\\\o "1-3" \\\\h \\\\z \\\\u')
    paragraph._p.append(fld)


def main(output_path: str = "final.docx") -> Path:
    doc = Document()

    # Skip cover page and TOC per request

    # Section 5.2
    if SECTION5_2_DRA_PATH and REPORTING_PERIOD:
        s5_2 = accumulate_section5_2(SECTION5_2_DRA_PATH, REPORTING_PERIOD, medname=MEDICINE or 'Product')
        write_section_5_2(
            doc,
            nstudies=s5_2.nstudies,
            medname=s5_2.medname,
            reporting_period=s5_2.reporting_period,
            total_subjects=s5_2.total_subjects,
            gender_text=s5_2.gender_text,
            age_text=s5_2.age_text,
            race_text=s5_2.race_text,
            table_structure=s5_2.table_structure,
        )

    # Section 5.3
    if SECTION5_3_DOCX_PATH:
        s5_3 = accumulate_section5_3(
            docx_path=SECTION5_3_DOCX_PATH,
            ddd_excel_path=DDD_EXCEL_PATH or None,
            country=COUNTRY,
            medicine=MEDICINE or 'Product',
            place=PLACE,
            date=DATE or '',
        )
        if s5_3.results is None or s5_3.results.combined_total == 0:
            generate_fallback_doc(doc, s5_3.medicine)
        else:
            generate_exposure_document(
                doc,
                country_table=s5_3.results.country_table,
                non_country_table=s5_3.results.non_country_table,
                combined_total=s5_3.results.combined_total,
                country_total=s5_3.results.country_total,
                non_country_total=s5_3.results.non_country_total,
                medicine=s5_3.medicine,
                place=s5_3.place,
                date=s5_3.date,
                country_name=s5_3.country_name,
            )

    # Section 6.3
    if CUM_EXCEL_PATH and CUM_RTF_PATH:
        s6_3 = accumulate_section6_3(
            cumulative_excel=CUM_EXCEL_PATH,
            cumulative_rtf=CUM_RTF_PATH,
            interval_excel=INT_EXCEL_PATH or None,
            interval_rtf=INT_RTF_PATH or None,
        )
        write_section_6_3(doc, cumulative_text=s6_3.cumulative_text, interval_text=s6_3.interval_text)

    # Section 15
    if SECTION15_INPUT_PATH:
        s15 = accumulate_section15(SECTION15_INPUT_PATH)
        write_section_15(doc, table_section15=s15.table, closed_signal=s15.closed_signals)

    # Apply branding (header/footer) using provided helper
    gen = PSURGenerator(
        logo_path=LOGO_PATH,
        drug_name=DRUG_NAME or 'Product',
        reporting_period=REPORTING_PERIOD_COVER or '',
        version_status=VERSION_STATUS,
        version_date=VERSION_DATE,
        drug_code=DRUG_CODE,
    )
    gen.apply_branding_to_document(document=doc)

    out = Path(output_path)
    doc.save(str(out))
    return out


if __name__ == "__main__":
    main()

